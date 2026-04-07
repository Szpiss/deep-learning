from __future__ import annotations

import json
import os
import platform
import re
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import torch
import torchvision
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
RUNS_DIR = ROOT / "runs" / "vgg16"
CHECKPOINT_DIR = ROOT / "checkpoint" / "vgg16"
ARTIFACTS_DIR = ROOT / "artifacts"
FIGURES_DIR = ROOT / "figures"
REPORTS_DIR = ROOT / "reports"
OUTPUT_DOCX = REPORTS_DIR / "实验02_利用VGG16进行CIFAR100分类_实验报告.docx"


def latest_subdir(path: Path) -> Path:
    subdirs = [item for item in path.iterdir() if item.is_dir()]
    if not subdirs:
        raise FileNotFoundError(f"No subdirectories found in {path}")
    return max(subdirs, key=lambda item: item.stat().st_mtime)


def latest_best_weight(path: Path) -> Path:
    def epoch_num(file_path: Path) -> int:
        match = re.search(r"-(\d+)-best\.pth$", file_path.name)
        return int(match.group(1)) if match else -1

    weight_files = sorted(path.glob("*-best.pth"), key=epoch_num)
    if not weight_files:
        raise FileNotFoundError(f"No best weight file found in {path}")
    return weight_files[-1]


def parse_test_log(path: Path) -> dict[str, str]:
    result = {
        "top1_err": "未执行",
        "top5_err": "未执行",
        "params": "未执行",
    }
    if not path.exists():
        return result

    text = path.read_text(encoding="utf-8", errors="ignore")
    patterns = {
        "top1_err": r"Top 1 err:\s*tensor\(([0-9.]+)",
        "top5_err": r"Top 5 err:\s*tensor\(([0-9.]+)",
        "params": r"Parameter numbers:\s*([0-9]+)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if not match:
            continue
        for group in match.groups():
            if group:
                result[key] = group
                break
    return result


def build_figures(history: list[dict[str, float]]) -> tuple[Path, Path]:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)

    epochs = [item["epoch"] for item in history]
    train_loss = [item["train_loss"] for item in history]
    test_loss = [item["test_loss"] for item in history]
    accuracy = [item["accuracy"] for item in history]

    loss_path = FIGURES_DIR / "vgg16_loss_curve.png"
    acc_path = FIGURES_DIR / "vgg16_accuracy_curve.png"

    plt.figure(figsize=(8, 5))
    plt.plot(epochs, train_loss, marker="o", label="Train Loss")
    plt.plot(epochs, test_loss, marker="s", label="Test Loss")
    plt.xlabel("Epoch")
    plt.ylabel("Loss")
    plt.title("VGG16 on CIFAR-100 Loss Curve")
    plt.grid(True, linestyle="--", alpha=0.4)
    plt.legend()
    plt.tight_layout()
    plt.savefig(loss_path, dpi=200)
    plt.close()

    plt.figure(figsize=(8, 5))
    plt.plot(epochs, accuracy, marker="o", color="#1f77b4")
    plt.xlabel("Epoch")
    plt.ylabel("Accuracy")
    plt.title("VGG16 on CIFAR-100 Accuracy Curve")
    plt.grid(True, linestyle="--", alpha=0.4)
    plt.tight_layout()
    plt.savefig(acc_path, dpi=200)
    plt.close()

    return loss_path, acc_path


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    run.font.size = Pt(9)


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)


def main() -> None:
    latest_run = latest_subdir(RUNS_DIR)
    latest_ckpt = latest_subdir(CHECKPOINT_DIR)
    best_weight = latest_best_weight(latest_ckpt)
    history_path = latest_run / "history.json"
    test_log_path = ARTIFACTS_DIR / "vgg16_test_best.log"

    history = json.loads(history_path.read_text(encoding="utf-8"))
    best_record = max(history, key=lambda item: item["accuracy"])
    final_record = history[-1]
    loss_fig, acc_fig = build_figures(history)
    test_result = parse_test_log(test_log_path)

    from models.vgg import vgg16_bn

    model = vgg16_bn()
    parameter_count = sum(param.numel() for param in model.parameters())

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("深度学习实验报告")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("第2题：利用VGG-16进行CIFAR-100分类")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)

    document.add_paragraph(
        "说明：本报告依据“深度学习实验报告模板”和“实验02 利用VGG-16进行CIFAR-100分类”的题目要求整理，"
        "已补充实验环境、运行命令、实验结果与分析内容。学号姓名可在最终提交前按老师要求修改文件名。"
    )

    add_section_heading(document, "题目要求")
    document.add_paragraph(
        "使用 VGG-16 网络对 CIFAR-100 数据集进行图像分类，熟悉数据加载、模型训练、模型测试的完整流程。"
        "在本机环境下完成模型训练与测试，并根据结果分析模型表现。"
    )
    document.add_paragraph(
        "代码来源：pytorch-cifar100（https://github.com/weiaicunzai/pytorch-cifar100），"
        "本实验在原项目基础上补充了 Apple Silicon MPS 设备支持、自定义训练轮数参数和训练历史保存功能，"
        "便于在当前机器上完成 10 轮实验并生成报告图表。"
    )

    add_section_heading(document, "实验环境")
    document.add_paragraph(f"操作系统：{platform.platform()}")
    document.add_paragraph(f"Python 版本：{platform.python_version()}")
    document.add_paragraph(f"PyTorch 版本：{torch.__version__}")
    document.add_paragraph(f"Torchvision 版本：{torchvision.__version__}")
    document.add_paragraph("训练设备：Apple Silicon MPS（Metal Performance Shaders）")
    document.add_paragraph("数据集：CIFAR-100，训练集 50000 张图片，测试集 10000 张图片，共 100 个类别。")

    add_section_heading(document, "实验原理")
    document.add_paragraph(
        "VGG-16 是一种经典卷积神经网络结构，其特点是使用多个 3x3 卷积层堆叠来逐步提取图像特征，"
        "并通过池化层降低空间分辨率，最后使用全连接层完成分类。"
    )
    document.add_paragraph(
        "本实验中的 VGG16_BN 结构包含 13 个卷积层和 3 个全连接层，并加入 Batch Normalization。"
        "输入图像大小为 32x32，输出类别数为 100。模型参数量约为 "
        f"{parameter_count:,}。"
    )

    add_section_heading(document, "实现代码")
    document.add_paragraph("1. VGG-16 网络核心定义（models/vgg.py）")
    add_code_block(
        document,
        "def vgg16_bn():\n"
        "    return VGG(make_layers(cfg['D'], batch_norm=True))\n\n"
        "self.classifier = nn.Sequential(\n"
        "    nn.Linear(512, 4096), nn.ReLU(inplace=True), nn.Dropout(),\n"
        "    nn.Linear(4096, 4096), nn.ReLU(inplace=True), nn.Dropout(),\n"
        "    nn.Linear(4096, num_class)\n"
        ")"
    )
    document.add_paragraph("2. 训练脚本核心流程（train.py）")
    add_code_block(
        document,
        "cifar100_training_loader = get_training_dataloader(...)\n"
        "cifar100_test_loader = get_test_dataloader(...)\n"
        "optimizer = optim.SGD(net.parameters(), lr=0.1, momentum=0.9, weight_decay=5e-4)\n"
        "for epoch in range(1, args.epochs + 1):\n"
        "    train_loss = train(epoch)\n"
        "    test_loss, acc = eval_training(epoch)\n"
        "    if best_acc < acc:\n"
        "        torch.save(net.state_dict(), best_weights_path)"
    )
    document.add_paragraph("3. 测试脚本核心流程（test.py）")
    add_code_block(
        document,
        "net.load_state_dict(torch.load(args.weights, map_location=device))\n"
        "output = net(image)\n"
        "_, pred = output.topk(5, 1, largest=True, sorted=True)\n"
        "correct_1 += correct[:, :1].sum()\n"
        "correct_5 += correct[:, :5].sum()"
    )
    document.add_paragraph("4. 本次实际运行命令")
    add_code_block(
        document,
        "./.venv/bin/python train.py -net vgg16 -device mps -epochs 10 -b 128 -num-workers 0\n"
        "./.venv/bin/python test.py -net vgg16 -device mps -weights <best_weight_path> -b 128 -num-workers 0"
    )

    add_section_heading(document, "输入输出")
    document.add_paragraph(
        f"训练历史文件：{history_path.relative_to(ROOT)}\n"
        f"最佳权重文件：{best_weight.relative_to(ROOT)}"
    )
    document.add_paragraph(
        f"最终第 {final_record['epoch']} 轮："
        f"train_loss={final_record['train_loss']:.4f}，"
        f"test_loss={final_record['test_loss']:.4f}，"
        f"accuracy={final_record['accuracy']:.4f}"
    )
    document.add_paragraph(
        f"最佳结果出现在第 {best_record['epoch']} 轮："
        f"accuracy={best_record['accuracy']:.4f}"
    )
    document.add_paragraph(
        "测试脚本输出："
        f"Top-1 error={test_result['top1_err']}，"
        f"Top-5 error={test_result['top5_err']}，"
        f"Parameter numbers={test_result['params']}"
    )

    document.add_paragraph("图1 训练与测试损失曲线：")
    document.add_picture(str(loss_fig), width=Inches(6.2))
    document.add_paragraph("图2 测试准确率曲线：")
    document.add_picture(str(acc_fig), width=Inches(6.2))

    document.add_paragraph("图3 需补充截图：训练命令运行终端截图。")
    document.add_paragraph("建议截取内容：命令行、若干 epoch 日志、单轮结束后的 Accuracy 输出。")
    document.add_paragraph("图4 需补充截图：测试命令与 Top-1/Top-5 error 终端截图。")
    document.add_paragraph("建议截取内容：测试命令、权重文件路径、Top 1 err、Top 5 err、Parameter numbers。")

    add_section_heading(document, "结果分析")
    document.add_paragraph(
        "从训练曲线可以看出，随着 epoch 增加，训练损失与测试损失总体呈下降趋势，测试准确率逐步提高，"
        "说明 VGG-16 能够在 CIFAR-100 数据集上逐步学习到有效特征。"
    )
    document.add_paragraph(
        "由于本实验仅按课程要求训练 10 轮，尚未达到 README 中 200 轮训练时的最佳性能，因此当前结果更多体现"
        "的是模型训练流程的正确性和早期收敛趋势。若继续训练更多 epoch，并配合学习率衰减，准确率通常还能进一步提升。"
    )
    document.add_paragraph(
        "本机为 Apple Silicon 平台，因此实验采用 MPS 作为硬件加速后端。与纯 CPU 相比，MPS 能明显缩短训练时间，"
        "更适合本次课程实验。"
    )

    add_section_heading(document, "结论")
    document.add_paragraph(
        "本实验完成了基于 VGG-16 的 CIFAR-100 图像分类任务，实现了数据加载、模型训练、权重保存、模型测试和结果分析的完整流程。"
        "实验结果表明，VGG-16 在 CIFAR-100 上能够逐步提升分类性能，验证了深度卷积神经网络在图像分类任务中的有效性。"
    )

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Report written to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
