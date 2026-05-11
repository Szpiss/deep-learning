from __future__ import annotations

import json
import platform
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SAVE_DIR = ROOT / "save"
IMG_DIR = ROOT / "img"
ARTIFACTS_DIR = ROOT / "artifacts"
OUTPUT_DOCX = ROOT / "reports" / "实验04_基于LSTM实现春联上联对下联_实验报告.docx"
TEMPLATE_DOCX = Path("/Users/cuing/dl_xor_lab/深度学习实验报告模板.docx")


def add_heading(document: Document, text: str, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    run.font.size = Pt(9)


def parse_test_log(path: Path) -> tuple[str, list[tuple[str, str]]]:
    text = path.read_text(encoding="utf-8")
    loss_match = re.search(r"Test Loss:\s*([0-9.]+)", text)
    loss_value = loss_match.group(1) if loss_match else "未知"
    examples: list[tuple[str, str]] = []
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for index, line in enumerate(lines):
        if line.startswith("上联：") and index + 1 < len(lines) and lines[index + 1].startswith("下联："):
            examples.append((line.replace("上联：", ""), lines[index + 1].replace("下联：", "")))
    return loss_value, examples


def main() -> None:
    history = json.loads((SAVE_DIR / "history.json").read_text(encoding="utf-8"))
    test_loss, examples = parse_test_log(ARTIFACTS_DIR / "lstm_test.log")
    best_record = min(history, key=lambda item: item["valid_loss"])
    final_record = history[-1]

    if TEMPLATE_DOCX.exists():
        document = Document(TEMPLATE_DOCX)
        for _ in range(len(document.paragraphs)):
            p = document.paragraphs[0]
            p._element.getparent().remove(p._element)
    else:
        document = Document()

    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("深度学习实验报告")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("第4题：基于 LSTM 实现春联上联对下联")
    run.bold = True
    run.font.size = Pt(14)

    add_heading(document, "题目要求")
    document.add_paragraph(
        "本实验要求基于 LSTM 模型实现春联上联对下联任务，熟悉数据加载、模型训练和模型测试流程。"
        "代码来源于 ChineseCouplet 项目，本次实验在原始代码基础上补充了参数化训练、MPS 设备支持、"
        "测试脚本和实验日志保存，便于在当前机器上完成 10 轮训练并整理实验结果。"
    )
    document.add_paragraph(
        "数据集由训练集和测试集组成，其中 train_in.txt 与 train_out.txt 分别保存上联和下联训练样本，"
        "test_in.txt 与 test_out.txt 保存测试样本。项目会在训练时自动构建词表，并额外引入 UNK 处理未登录词。"
    )

    add_heading(document, "实验环境")
    document.add_paragraph(f"操作系统：{platform.platform()}")
    document.add_paragraph("Python 版本：3.12.13")
    document.add_paragraph("PyTorch 版本：2.11.0")
    document.add_paragraph("训练设备：Apple Silicon MPS")
    document.add_paragraph("训练轮数：10 轮")

    add_heading(document, "实现代码")
    document.add_paragraph("1. 模型结构核心代码")
    add_code(
        document,
        "class LSTM(nn.Module):\n"
        "    def __init__(self, vocab_size, embedding_dim, hidden_dim, num_layers):\n"
        "        super().__init__()\n"
        "        self.embeddings = nn.Embedding(vocab_size, embedding_dim)\n"
        "        self.lstm = nn.LSTM(embedding_dim, hidden_dim, num_layers, dropout=0.5)\n"
        "        self.linear = nn.Linear(hidden_dim, vocab_size)\n"
        "\n"
        "    def forward(self, x):\n"
        "        embeds = self.embeddings(x)\n"
        "        output, _ = self.lstm(embeds)\n"
        "        return self.linear(output.reshape(output.size(0) * output.size(1), -1))"
    )
    document.add_paragraph("2. 训练流程核心代码")
    add_code(
        document,
        "train_dict, vocab_size, idx2word, word2idx, max_len = load_data(...)\n"
        "model = LSTM(vocab_size=vocab_size, hidden_dim=512, embedding_dim=128, num_layers=3).to(device)\n"
        "optimizer = optim.Adam(model.parameters(), lr=1e-3)\n"
        "criterion = nn.CrossEntropyLoss()\n"
        "for epoch in range(1, args.epochs + 1):\n"
        "    ...\n"
        "    if valid_loss_mean < best_val_loss:\n"
        "        torch.save(best_state_dict, 'save/best_model.pt')"
    )
    document.add_paragraph("3. 测试流程核心代码")
    add_code(
        document,
        "checkpoint = torch.load('save/best_model.pt', map_location=device)\n"
        "model.load_state_dict(checkpoint['model_state_dict'])\n"
        "for x, y in data_generator(test_dict, batch_size=args.batch_size, max_len=max_len):\n"
        "    output = model(x_tensor)\n"
        "    loss = criterion(output, y_tensor.view(-1))"
    )
    document.add_paragraph("4. 本次实际运行命令")
    add_code(
        document,
        "./.venv/bin/python train.py --epochs 10 --device mps\n"
        "./.venv/bin/python test.py --device mps"
    )

    add_heading(document, "输入输出")
    document.add_paragraph(
        f"最佳验证损失出现在第 {best_record['epoch']} 轮，为 {best_record['valid_loss']:.6f}。"
    )
    document.add_paragraph(
        f"第 10 轮训练结束时，训练损失为 {final_record['train_loss']:.6f}，"
        f"验证损失为 {final_record['valid_loss']:.6f}，测试损失为 {final_record['test_loss']:.6f}。"
    )
    document.add_paragraph(f"测试脚本单独运行得到的测试损失为 {test_loss}。")

    document.add_paragraph("图1 LSTM 网络结构示意图：")
    document.add_picture(str(IMG_DIR / "img2.png"), width=Inches(5.8))

    document.add_paragraph("图2 训练损失、验证损失和测试损失变化曲线：")
    document.add_picture(str(SAVE_DIR / "loss.png"), width=Inches(6.0))

    document.add_paragraph("模型测试样例：")
    table = document.add_table(rows=1, cols=2)
    header_cells = table.rows[0].cells
    header_cells[0].text = "上联"
    header_cells[1].text = "生成下联"
    for upper, lower in examples:
        row_cells = table.add_row().cells
        row_cells[0].text = upper
        row_cells[1].text = lower

    document.add_paragraph("训练日志摘要：")
    add_code(
        document,
        (ARTIFACTS_DIR / "lstm_train_10epochs.log").read_text(encoding="utf-8")
    )
    document.add_paragraph("测试日志摘要：")
    add_code(
        document,
        (ARTIFACTS_DIR / "lstm_test.log").read_text(encoding="utf-8")
    )

    add_heading(document, "结果分析")
    document.add_paragraph(
        "从 10 轮训练结果来看，训练损失、验证损失和测试损失总体呈下降趋势，说明 LSTM 模型能够逐步学习春联上下联之间的序列映射关系。"
        "虽然仅训练了 10 轮，模型生成的下联在语义和工整性上仍然有限，但已经具备了基本的字面续写能力。"
    )
    document.add_paragraph(
        "从测试样例可以看到，模型能够输出长度匹配的汉字序列，但下联内容仍存在重复字较多、语义连贯性不足的问题。"
        "这与训练轮数较少、模型结构相对基础以及没有使用注意力机制等因素有关。若增加训练轮数或改进为 Seq2Seq + Attention，结果通常会更自然。"
    )

    add_heading(document, "结论")
    document.add_paragraph(
        "本实验完成了基于 LSTM 的春联对下联生成任务，实现了数据读取、词表构建、模型训练、模型测试和实验报告整理的完整流程。"
        "实验结果表明，LSTM 能够用于中文序列生成任务，并能在春联自动对联场景中学习到一定的上下文映射规律。"
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Report written to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
